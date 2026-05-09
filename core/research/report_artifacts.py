from __future__ import annotations

import re
import textwrap
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape


SUPPORTED_DERIVED_FORMATS = ("pdf", "docx")

_FORMAT_ALIASES = {
    "pdf": "pdf",
    "application/pdf": "pdf",
    "doc": "docx",
    "docx": "docx",
    "word": "docx",
    "msword": "docx",
    "openxml": "docx",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}
_MARKDOWN_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_MARKDOWN_STYLE_RE = re.compile(r"[*_`]+")


def normalize_derived_formats(*values: Any) -> list[str]:
    formats: list[str] = []

    def add(raw: Any) -> None:
        if raw is None:
            return
        if isinstance(raw, dict):
            for key, enabled in raw.items():
                if enabled:
                    add(key)
            return
        if isinstance(raw, (list, tuple, set)):
            for item in raw:
                add(item)
            return
        text = str(raw or "").strip().lower()
        if not text:
            return
        for part in re.split(r"[\s,;|+]+", text):
            candidate = part.strip().lower()
            if not candidate or candidate in {"md", "markdown", "text"}:
                continue
            if candidate in {"all", "office"}:
                for fmt in SUPPORTED_DERIVED_FORMATS:
                    if fmt not in formats:
                        formats.append(fmt)
                continue
            normalized = _FORMAT_ALIASES.get(candidate)
            if normalized and normalized not in formats:
                formats.append(normalized)

    for value in values:
        add(value)
    return formats


def requested_research_report_formats(task) -> list[str]:
    policy = dict(getattr(task, "source_policy", {}) or {})
    formats = normalize_derived_formats(
        policy.get("derived_formats"),
        policy.get("artifact_formats"),
        policy.get("report_formats"),
        getattr(task, "output_format", ""),
    )
    bool_flags = (
        ("pdf", ("pdf", "include_pdf", "derive_pdf")),
        ("docx", ("docx", "include_docx", "derive_docx")),
    )
    for fmt, keys in bool_flags:
        if fmt in formats:
            continue
        if any(_truthy(policy.get(key)) for key in keys):
            formats.append(fmt)
    return formats


def create_research_report_derivatives(
    artifact_service,
    *,
    task,
    report_markdown: str,
    source_artifact,
    citation_validation: dict[str, Any] | None = None,
    runner: str = "core.research_execution.v1",
    requested_formats: list[str] | None = None,
) -> list[dict[str, Any]]:
    formats = requested_formats if requested_formats is not None else requested_research_report_formats(task)
    normalized_formats = normalize_derived_formats(formats)
    if not normalized_formats:
        return []

    source_artifact_id = str(getattr(source_artifact, "artifact_id", "") or "")
    source_filename = str(getattr(source_artifact, "filename", "") or "research-report.md")
    filename_stem = Path(source_filename).stem or str(getattr(task, "research_task_id", "") or "research-report")
    title = str(getattr(task, "topic", "") or filename_stem)
    created: list[dict[str, Any]] = []
    for fmt in normalized_formats:
        data, content_type, extension = build_report_derivative_bytes(
            fmt,
            report_markdown=report_markdown,
            title=title,
        )
        artifact = artifact_service.create_bytes_artifact(
            principal_id=getattr(task, "principal_id", None),
            project_id=getattr(task, "project_id", None),
            thread_id=getattr(task, "thread_id", None),
            data=data,
            filename=f"{filename_stem}.{extension}",
            artifact_type="research_report_derivative",
            content_type=content_type,
            metadata={
                "research_task_id": str(getattr(task, "research_task_id", "") or ""),
                "runner": runner,
                "derived_format": fmt,
                "source_artifact_id": source_artifact_id,
                "citation_validation": dict(citation_validation or {}),
            },
        )
        created.append(
            {
                "format": fmt,
                "artifact_id": artifact.artifact_id,
                "filename": artifact.filename,
                "content_type": artifact.content_type,
                "byte_size": int(artifact.byte_size or 0),
                "checksum": artifact.checksum,
                "download_url": f"/runtime/artifacts/{artifact.artifact_id}/download",
                "source_artifact_id": source_artifact_id,
            }
        )
    return created


def build_report_derivative_bytes(format_name: str, *, report_markdown: str, title: str = "") -> tuple[bytes, str, str]:
    normalized = normalize_derived_formats(format_name)
    if not normalized:
        raise ValueError(f"unsupported derived report format: {format_name}")
    fmt = normalized[0]
    if fmt == "pdf":
        return (
            build_text_pdf_bytes(report_markdown, title=title),
            "application/pdf",
            "pdf",
        )
    if fmt == "docx":
        return (
            build_docx_bytes(report_markdown, title=title),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        )
    raise ValueError(f"unsupported derived report format: {format_name}")


def build_docx_bytes(report_markdown: str, *, title: str = "") -> bytes:
    try:
        from docx import Document
    except ImportError:
        return _build_minimal_docx_bytes(report_markdown)

    document = Document()
    if title:
        document.core_properties.title = title
    for raw_line in str(report_markdown or "").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            document.add_paragraph()
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            document.add_heading(_clean_markdown_inline(heading.group(2)), level=min(len(heading.group(1)), 4))
            continue
        if stripped.startswith("- "):
            document.add_paragraph(_clean_markdown_inline(stripped[2:]), style="List Bullet")
            continue
        document.add_paragraph(_clean_markdown_inline(stripped))
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_text_pdf_bytes(report_markdown: str, *, title: str = "") -> bytes:
    lines = _wrapped_plain_lines(report_markdown, title=title, width=86)
    pages = [lines[index : index + 46] for index in range(0, len(lines), 46)] or [[""]]

    objects: dict[int, bytes] = {}
    page_ids: list[int] = []
    next_object_id = 5
    for page_lines in pages:
        page_id = next_object_id
        content_id = next_object_id + 1
        next_object_id += 2
        page_ids.append(page_id)
        stream = _pdf_page_stream(page_lines)
        objects[content_id] = (
            f"<< /Length {len(stream)} >>\nstream\n".encode("ascii")
            + stream
            + b"\nendstream"
        )
        objects[page_id] = (
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 3 0 R >> >> /Contents {content_id} 0 R >>"
        ).encode("ascii")

    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    objects[1] = b"<< /Type /Catalog /Pages 2 0 R >>"
    objects[2] = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>".encode("ascii")
    objects[3] = (
        b"<< /Type /Font /Subtype /Type0 /BaseFont /STSong-Light "
        b"/Encoding /UniGB-UCS2-H /DescendantFonts [4 0 R] >>"
    )
    objects[4] = (
        b"<< /Type /Font /Subtype /CIDFontType0 /BaseFont /STSong-Light "
        b"/CIDSystemInfo << /Registry (Adobe) /Ordering (GB1) /Supplement 2 >> /DW 1000 >>"
    )

    output = BytesIO()
    output.write(b"%PDF-1.4\n%\xE2\xE3\xCF\xD3\n")
    offsets = [0]
    for object_id in range(1, max(objects) + 1):
        offsets.append(output.tell())
        output.write(f"{object_id} 0 obj\n".encode("ascii"))
        output.write(objects[object_id])
        output.write(b"\nendobj\n")
    xref_offset = output.tell()
    output.write(f"xref\n0 {len(offsets)}\n".encode("ascii"))
    output.write(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.write(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.write(
        f"trailer\n<< /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    return output.getvalue()


def _truthy(value: Any) -> bool:
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "on", "enabled"}
    return bool(value)


def _plain_markdown_lines(markdown: str, *, title: str = "") -> list[str]:
    lines: list[str] = []
    if title:
        lines.extend([title.strip(), ""])
    for raw_line in str(markdown or "").splitlines():
        stripped = raw_line.strip()
        if not stripped:
            lines.append("")
            continue
        heading = re.match(r"^#{1,6}\s+(.+)$", stripped)
        if heading:
            lines.append(_clean_markdown_inline(heading.group(1)))
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            lines.append("- " + _clean_markdown_inline(bullet.group(1)))
            continue
        lines.append(_clean_markdown_inline(stripped))
    return lines or [""]


def _wrapped_plain_lines(markdown: str, *, title: str = "", width: int = 86) -> list[str]:
    wrapped: list[str] = []
    for line in _plain_markdown_lines(markdown, title=title):
        if not line:
            wrapped.append("")
            continue
        wrapped.extend(
            textwrap.wrap(
                line,
                width=max(20, int(width or 86)),
                break_long_words=True,
                break_on_hyphens=False,
                replace_whitespace=False,
            )
            or [""]
        )
    return wrapped


def _clean_markdown_inline(text: str) -> str:
    value = _MARKDOWN_LINK_RE.sub(r"\1 (\2)", str(text or ""))
    value = _MARKDOWN_STYLE_RE.sub("", value)
    return " ".join(value.split())


def _pdf_page_stream(lines: list[str]) -> bytes:
    commands = ["BT", "/F1 10 Tf", "14 TL", "72 740 Td"]
    for index, line in enumerate(lines):
        if index > 0:
            commands.append("T*")
        text = _pdf_text_hex(line)
        commands.append(f"<{text}> Tj")
    commands.append("ET")
    return "\n".join(commands).encode("ascii")


def _pdf_text_hex(text: str) -> str:
    cleaned = "".join(ch if ch >= " " or ch == "\t" else " " for ch in str(text or ""))
    return cleaned.encode("utf-16-be", errors="replace").hex().upper()


def _build_minimal_docx_bytes(report_markdown: str) -> bytes:
    paragraphs = []
    for line in _plain_markdown_lines(report_markdown):
        paragraphs.append(
            "<w:p><w:r><w:t xml:space=\"preserve\">"
            + xml_escape(line)
            + "</w:t></w:r></w:p>"
        )
    document_xml = (
        "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
        "<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
        "<w:body>"
        + "".join(paragraphs)
        + "<w:sectPr><w:pgSz w:w=\"12240\" w:h=\"15840\"/><w:pgMar w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\"/></w:sectPr>"
        "</w:body></w:document>"
    )
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            "<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
            "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
            "<Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>"
            "<Default Extension=\"xml\" ContentType=\"application/xml\"/>"
            "<Override PartName=\"/word/document.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>"
            "</Types>",
        )
        archive.writestr(
            "_rels/.rels",
            "<?xml version=\"1.0\" encoding=\"UTF-8\"?>"
            "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
            "<Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" Target=\"word/document.xml\"/>"
            "</Relationships>",
        )
        archive.writestr("word/document.xml", document_xml)
    return buffer.getvalue()
